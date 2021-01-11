from docx import Document
import pandas as pd 

index_excel = os.path.join(os.path.dirname(__file__), 'index.xlsx')
export_doc = os.path.join(os.path.dirname(__file__), '隔页.doc')

df = pd.read_excel(index_excel)
doc = Document()
for i in range(len(df)):
    if ("参见" in str(df.iloc[i,2])) or ("不适用" in str(df.iloc[i,2])):
        index = '{} {}，{}'.format(df.iloc[i,0], df.iloc[i,1], df.iloc[i,2])
        doc.add_paragraph(index)    
    else:
        index = '{} {}'.format(df.iloc[i,0], df.iloc[i,1])
        doc.add_paragraph(index)

    if i < len(df)-1:
        doc.add_page_break()

doc.save(export_doc)